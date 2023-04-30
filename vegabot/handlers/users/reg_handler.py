from aiogram import types
from aiogram.dispatcher import FSMContext
from docx import Document
from docx.shared import Inches, Pt
import asyncio
import os
from datetime import datetime
from loader import dp, db, bot
from messages.btn_text.btn_text import *
from messages.funcs.check_user import check_user
from messages.funcs.get_keyboard_default import get_markup_default, get_markup_default_phone
from messages.funcs.get_keyboard_inline import get_markup_inline
from messages.funcs.get_language import get_language
from messages.message_txt.message_text import *
from states.full_info import Full_info_user
from data.config import channel_id
from keyboards.inline.know_lang import get_kb, get_markup
from keyboards.default.jobs_button import get_jobs_button

know_lan = []


@dp.message_handler(text=["📝 Ariza qoldirish", "📝 Оставить заявку"])
async def get_name(message: types.Message):
    if await check_user(user_id=message.from_user.id) == False:
        await db.add_user(telegram_id=message.from_user.id,
                          language='uz',
                          fullname=None,
                          date=None,
                          day=None,
                          phone=None,
                          addphone=None,
                          malumot=None,
                          manzil=None,
                          oilaviy_holat=None,
                          sud_holat=None,
                          company=None,
                          new_company=None,
                          last_summa=None,new_summa=None,time_job=None, know_lang=None,know_lang_degree=None,
                          abaut_as=None)

    language = await get_language(message.from_user.id)
    if language is None:
        language = 'uz'
    # know_lang_txt_btn[language].clear()
    # for value in know_lang_txt_btn1[language]:
    #     know_lang_txt_btn[language].append(value)
    await message.answer(name_txt[language], reply_markup=types.ReplyKeyboardRemove())
    await Full_info_user.fullname.set()


@dp.message_handler(content_types=['text'], state=Full_info_user.fullname)
async def get_name(message: types.Message, state: FSMContext):
    language = await get_language(message.from_user.id)
    await state.update_data(fullname=message.text)
    await message.answer(date_txt[language])
    await Full_info_user.date.set()


@dp.message_handler(content_types=['text'], state=Full_info_user.date)
async def get_name(message: types.Message, state: FSMContext):
    language = await get_language(message.from_user.id)
    await state.update_data(date=message.text)
    await message.answer(phone_txt[language], reply_markup=await get_markup_default_phone(language, phone_btn_txt))
    await Full_info_user.phone.set()


@dp.message_handler(content_types=['contact'], state=Full_info_user.phone)
async def get_name(message: types.Message, state: FSMContext):
    language = await get_language(message.from_user.id)
    await state.update_data(phone=message.contact.phone_number)
    await message.answer(addphone_txt[language], reply_markup=types.ReplyKeyboardRemove())
    await Full_info_user.addphone.set()


@dp.message_handler(content_types=['text'], state=Full_info_user.addphone)
async def get_name(message: types.Message, state: FSMContext):
    language = await get_language(message.from_user.id)
    if message.text.startswith('+998') and len(message.text) == 13:
        await state.update_data(addphone=message.text)
        await message.answer(malumot_txt[language], reply_markup=await get_markup_default(language, malumot_txt_btn))
        await Full_info_user.malumot.set()
    else:
        await message.answer(phone_error_txt[language])
        await Full_info_user.addphone.set()


@dp.message_handler(content_types=['text'], state=Full_info_user.malumot)
async def get_name(message: types.Message, state: FSMContext):
    language = await get_language(message.from_user.id)
    await state.update_data(malumot=message.text)
    await message.answer(manzil_txt[language], reply_markup=types.ReplyKeyboardRemove())
    await Full_info_user.manzil.set()


@dp.message_handler(content_types=['text'], state=Full_info_user.manzil)
async def get_name(message: types.Message, state: FSMContext):
    language = await get_language(message.from_user.id)
    await state.update_data(manzil=message.text)
    await message.answer(oilaviy_holat_txt[language],
                         reply_markup=await get_markup_default(language, oilaviy_holat_txt_btn))
    await Full_info_user.oilaviy_holat.set()


@dp.message_handler(content_types=['text'], state=Full_info_user.oilaviy_holat)
async def get_name(message: types.Message, state: FSMContext):
    language = await get_language(message.from_user.id)
    await state.update_data(oilaviy_holat=message.text)
    await message.answer(sud_holat_txt[language],
                         reply_markup=await get_markup_default(language, sud_holat_txt_btn))
    await Full_info_user.sud_holat.set()


@dp.message_handler(content_types=['text'], state=Full_info_user.sud_holat)
async def get_name(message: types.Message, state: FSMContext):
    language = await get_language(message.from_user.id)
    await state.update_data(sud_holat=message.text)
    await message.answer(qiziqishlari[language], reply_markup=types.ReplyKeyboardRemove())
    await Full_info_user.qiziqish.set()


@dp.message_handler(content_types=['text'], state=Full_info_user.qiziqish)
async def get_name(message: types.Message, state: FSMContext):
    language = await get_language(message.from_user.id)
    await state.update_data(qiziqish=message.text)
    await message.answer(chidamli[language])
    await Full_info_user.chidam.set()


@dp.message_handler(content_types=['text'], state=Full_info_user.chidam)
async def get_name(message: types.Message, state: FSMContext):
    language = await get_language(message.from_user.id)
    await state.update_data(chidam=message.text)
    await message.answer(rasm_txt[language])
    await Full_info_user.photo.set()


@dp.message_handler(content_types=['photo'], state=Full_info_user.photo)
async def get_name(message: types.Message, state: FSMContext):
    language = await get_language(message.from_user.id)
    await message.photo[-1].download(destination_file=f'handlers/users/image_{message.from_user.id}.jpg')
    await message.photo[-1].download(destination_file=f'image_{message.from_user.id}.jpg')
    await message.answer(company_txt[language])
    await Full_info_user.company.set()


@dp.message_handler(content_types=['text'], state=Full_info_user.company)
async def get_name(message: types.Message, state: FSMContext):
    language = await get_language(message.from_user.id)
    await state.update_data(company=message.text)
    
    await message.answer(newcompany_txt[language], reply_markup= await get_jobs_button(language))
    await Full_info_user.newcompany.set()


@dp.message_handler(content_types=['text'], state=Full_info_user.newcompany)
async def get_name(message: types.Message, state: FSMContext):
    language = await get_language(message.from_user.id)
    await state.update_data(newcompany=message.text)
    await message.answer(last_summa_txt[language], reply_markup=types.ReplyKeyboardRemove())
    await Full_info_user.last_summa.set()


@dp.message_handler(content_types=['text'], state=Full_info_user.last_summa)
async def get_name(message: types.Message, state: FSMContext):
    language = await get_language(message.from_user.id)
    await state.update_data(last_summa=message.text)
    await message.answer(new_summa_txt[language], reply_markup=types.ReplyKeyboardRemove())
    await Full_info_user.new_summa.set()


@dp.message_handler(content_types=['text'], state=Full_info_user.new_summa)
async def get_name(message: types.Message, state: FSMContext):
    language = await get_language(message.from_user.id)
    await state.update_data(new_summa=message.text)
    await message.answer(time_job_txt[language], reply_markup=await get_markup_default(language, time_job_txt_btn))
    await Full_info_user.time_job.set()


@dp.message_handler(content_types=['text'], state=Full_info_user.time_job)
async def get_name(message: types.Message, state: FSMContext):
    language = await get_language(message.from_user.id)
    await state.update_data(time_jobs=message.text)
    await message.answer(know_lang_change_txt[language], reply_markup=types.ReplyKeyboardRemove())
    await message.answer(know_lang_txt[language], reply_markup=await get_markup(know_lang_txt_btn[language]))
    await Full_info_user.know_lang.set()


@dp.callback_query_handler(state=Full_info_user.know_lang)
async def get_name(call: types.CallbackQuery, state: FSMContext):
    language = await get_language(call.from_user.id)
    if call.data.split('_')[0] in ["➡ Davom etish", "➡ Продолжить"]:
        await state.update_data(know_lang=set(know_lan))
        await call.message.answer(know_lang_degree_txt[language],
                                  reply_markup=await get_markup_default(language, know_lang_degree_txt_btn))
        await Full_info_user.know_lang_degree.set()
    else:
        if call.data.startswith('✅'):
            know_lan.remove(((call.data.split('_')[0]).split('✅')[1]).lstrip())
            await call.message.edit_reply_markup(reply_markup=await get_kb(data=call.data, language=language))
        else:
            know_lan.append(call.data.split('_')[0])
            await call.message.edit_reply_markup(reply_markup=await get_kb(data=call.data, language=language))
            await Full_info_user.know_lang.set()


@dp.message_handler(content_types=['text'], state=Full_info_user.know_lang_degree)
async def get_name(message: types.Message, state: FSMContext):
    language = await get_language(message.from_user.id)
    await state.update_data(know_lang_degree=message.text)
    await message.answer(nima_uchun_txt[language], reply_markup=types.ReplyKeyboardRemove())
    await Full_info_user.nima_uchun.set()


@dp.message_handler(content_types=['text'], state=Full_info_user.nima_uchun)
async def get_name(message: types.Message, state: FSMContext):
    language = await get_language(message.from_user.id)
    await state.update_data(nima_uchun=message.text)
    await message.answer(about_us_txt[language])
    await Full_info_user.about_us.set()


@dp.message_handler(content_types=['text'], state=Full_info_user.about_us)
async def get_name(message: types.Message, state: FSMContext):
    language = await get_language(message.from_user.id)
    await state.update_data(about_us=message.text)
    await message.answer(comfirm_finish[language], reply_markup=await get_markup_inline(language, confirm_txt_btn))
    await Full_info_user.comfirm.set()


@dp.callback_query_handler(lambda query: query.data == "✅ Tasdiqlash" or query.data == "✅ Подтвердить",
                           state=Full_info_user.comfirm)
async def get_name(query: types.CallbackQuery, state: FSMContext):
    language = await get_language(query.from_user.id)
    data = await state.get_data()

    name = data.get('fullname')
    dates = data.get('date')
    phone = data.get('phone')
    addphone = data.get('addphone')
    malumot = data.get('malumot')
    manzil = data.get('manzil')
    oilaviy_holat = data.get('oilaviy_holat')
    sud_holat = data.get('sud_holat')
    company = data.get('company')
    newcompany = data.get('newcompany')
    last_summa = data.get('last_summa')
    new_summa = data.get('new_summa')
    time_jobs = data.get('time_jobs')
    array = data.get('know_lang')
    know_lang_degree = data.get('know_lang_degree')
    nima_uchun = data.get('nima_uchun')
    about_us = data.get('about_us')
    know_lang = ','.join(array)

    document = Document()
    document.add_picture(f'image_{query.from_user.id}.jpg', width=Inches(1.25))
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times'
    font.size = Pt(13)

    if language == 'uz':
        document.add_paragraph("SHAXSIY MA'LUMOTLARI", style='Body Text').alignment = 1
        document.add_paragraph(f"Ism,familiyasi:  {name}", style='Body Text')
        document.add_paragraph(f"Tug'ilgan sanasi:  {dates}", style='Body Text')
        document.add_paragraph(f"Telefon raqami:  {phone}", style='Body Text')
        document.add_paragraph(f"Qo'shimcha tel:  {addphone}", style='Body Text')
        document.add_paragraph(f"Ma'lumoti:  {malumot}", style='Body Text')
        document.add_paragraph(f"Yashash manzili:  {manzil}", style='Body Text')
        document.add_paragraph(f"Oilaviy axvoli:  {oilaviy_holat}", style='Body Text')
        document.add_paragraph(f"Sudlanganmi:  {sud_holat}", style='Body Text')
        document.add_paragraph(f"ISH HAQIDA MA'LUMOT: ", style='Body Text').alignment = 1

        table = document.add_table(rows=9, cols=2)
        table.rows[0].cells[0].text = 'Qaysi tashkilotlarda va qaysi lavozimlarda ishlagan:'
        table.rows[1].cells[0].text = 'Bizning tashkilotda qaysi lavozimda ishlamoqchi:'
        table.rows[2].cells[0].text = "Oxirgi ish o'rnidagi oylik maoshi:"
        table.rows[3].cells[0].text = 'Bizda qancha miqdorli maoshga ishlamoqchi:'
        table.rows[4].cells[0].text = 'Bizning tashkilotda qancha muddat ishlamoqchi:'
        table.rows[5].cells[0].text = 'Qaysi tillarni biladi:'
        table.rows[6].cells[0].text = 'Bu tillarni qay darajada biladi:'
        table.rows[7].cells[0].text = 'Biz haqimizda qayerdan eshitgan:'
        table.rows[8].cells[0].text = 'Nima uchun aynan sizni ishga olishimiz kerak? '
        table.style = 'TableGrid'
        table.rows[0].cells[1].text = f"{company}"
        table.rows[1].cells[1].text = f"{newcompany}"
        table.rows[2].cells[1].text = f"{last_summa}"
        table.rows[3].cells[1].text = f"{new_summa}"
        table.rows[4].cells[1].text = f"{time_jobs}"
        table.rows[5].cells[1].text = f"{know_lang}"
        table.rows[6].cells[1].text = f"{know_lang_degree}"
        table.rows[7].cells[1].text = f"{about_us}"
        table.rows[8].cells[1].text = f"{nima_uchun}"

        document.save(f'{query.from_user.id}.docx')

        if types.InputFile(path_or_bytesio=f"{query.from_user.id}.docx"):
            await db.update_user_full(fullname=name,
                                      date=dates,
                                      phone=phone,
                                      addphone=addphone,
                                      malumot=malumot,
                                      manzil=manzil,
                                      oilaviy_holat=oilaviy_holat,
                                      sud_holat=sud_holat,
                                      day=datetime.now(),
                                      company=company,
                                      new_company=newcompany,
                                      last_summa=last_summa,
                                      new_summa=new_summa,
                                      time_job=time_jobs,
                                      know_lang=know_lang,
                                      know_lang_degree=know_lang_degree,
                                      abaut_as=about_us,
                                      telegram_id=query.from_user.id)

            await query.message.edit_text("Yuklanmoqda... 3")
            await asyncio.sleep(0.5)
            await query.message.edit_text("Yuklanmoqda... 2")
            await asyncio.sleep(0.5)
            await query.message.edit_text("Yuklanmoqda... 1")
            await query.message.delete()

            await bot.send_document(chat_id=int(channel_id),
                                    document=types.InputFile(path_or_bytesio=f"{query.from_user.id}.docx",
                                                             filename="Resume.docx"),
                                    caption=f"👤 Ism,familiyasi: {name}\n📞 Telefon raqami:  {phone}\n")

            await bot.send_document(chat_id=query.from_user.id,
                                    document=types.InputFile(path_or_bytesio=f"{query.from_user.id}.docx",
                                                             filename="Resume.docx"),
                                    reply_markup=await get_markup_default(language, reg_btn_txt))

            os.remove(f"{query.from_user.id}.docx")
            os.remove(f"image_{query.from_user.id}.jpg")
            os.remove(f"handlers/users/image_{query.from_user.id}.jpg")
            know_lan.clear()
            await state.finish()
        else:
            await query.message.answer("Xatolik yuz berdi qaytadan urining:")
            os.remove(f"{query.from_user.id}.docx")
            os.remove(f"image_{query.from_user.id}.jpg")
            os.remove(f"handlers/users/image_{query.from_user.id}.jpg")
            know_lan.clear()
            await state.finish()

    else:
        document.add_paragraph("ЛИЧНАЯ ИНФОРМАЦИЯ", style='Body Text').alignment = 1
        document.add_paragraph(f"Имя, Фамилия:  {name}", style='Body Text')
        document.add_paragraph(f"Твоя дата рождения:  {dates}", style='Body Text')
        document.add_paragraph(f"Ваш номер телефона:  {phone}", style='Body Text')
        document.add_paragraph(f"Дополнительный тел:  {addphone}", style='Body Text')
        document.add_paragraph(f"Ваша информация:  {malumot}", style='Body Text')
        document.add_paragraph(f"Адрес проживания::  {manzil}", style='Body Text')
        document.add_paragraph(f"Семейный статус:  {oilaviy_holat}", style='Body Text')
        document.add_paragraph(f"Вы были осуждены:  {sud_holat}", style='Body Text')
        document.add_paragraph(f"ИНФОРМАЦИЯ О РАБОТЕ: ", style='Body Text').alignment = 1

        table = document.add_table(rows=9, cols=2)
        table.rows[0].cells[0].text = 'В каких организациях и на каких должностях Вы работали:'
        table.rows[1].cells[0].text = 'На какой должности вы хотите работать в нашей организации:'
        table.rows[2].cells[0].text = "Ваша последняя зарплата:"
        table.rows[3].cells[0].text = 'Какую зарплату вы хотите у нас работать:'
        table.rows[4].cells[0].text = 'Как долго вы хотите работать в нашей организации:'
        table.rows[5].cells[0].text = 'Какие языки вы знаете:'
        table.rows[6].cells[0].text = 'Насколько хорошо вы знаете эти языки:'
        table.rows[7].cells[0].text = 'Где вы услышали о нас:'
        table.rows[8].cells[0].text = 'Почему именно мы должны нанять вас?'
        table.style = 'TableGrid'
        table.rows[0].cells[1].text = f"{company}"
        table.rows[1].cells[1].text = f"{newcompany}"
        table.rows[2].cells[1].text = f"{last_summa}"
        table.rows[3].cells[1].text = f"{new_summa}"
        table.rows[4].cells[1].text = f"{time_jobs}"
        table.rows[5].cells[1].text = f"{know_lang}"
        table.rows[6].cells[1].text = f"{know_lang_degree}"
        table.rows[7].cells[1].text = f"{about_us}"
        table.rows[8].cells[1].text = f"{nima_uchun}"

        document.save(f'{query.from_user.id}.docx')

        await state.finish()

        if types.InputFile(path_or_bytesio=f"{query.from_user.id}.docx"):
            await db.update_user_full(fullname=name,
                                      date=dates,
                                      phone=phone,
                                      addphone=addphone,
                                      malumot=malumot,
                                      manzil=manzil,
                                      oilaviy_holat=oilaviy_holat,
                                      sud_holat=sud_holat,
                                      day=datetime.now(),
                                      company=company,
                                      new_company=newcompany,
                                      last_summa=last_summa,
                                      new_summa=new_summa,
                                      time_job=time_jobs,
                                      know_lang=know_lang,
                                      know_lang_degree=know_lang_degree,
                                      abaut_as=about_us,
                                      telegram_id=query.from_user.id)

            await query.message.edit_text("Загрузка... 3")
            await asyncio.sleep(0.5)
            await query.message.edit_text("Загрузка... 2")
            await asyncio.sleep(0.5)
            await query.message.edit_text("Загрузка... 1")
            await query.message.delete()
            await bot.send_document(chat_id=int(channel_id),
                                    document=types.InputFile(path_or_bytesio=f"{query.from_user.id}.docx",
                                                             filename="Резюме.docx"),
                                    caption=f"👤 Имя, Фамилия: {name}\n📞 Hомер телефона:  {phone}\n📝 Oтделение: {newcompany}")

            await bot.send_document(chat_id=query.from_user.id,
                                    document=types.InputFile(path_or_bytesio=f"{query.from_user.id}.docx",
                                                             filename="Резюме.docx"),
                                    reply_markup=await get_markup_default(language, reg_btn_txt))
            os.remove(f"{query.from_user.id}.docx")
            os.remove(f"image_{query.from_user.id}.jpg")
            os.remove(f"handlers/users/image_{query.from_user.id}.jpg")
            know_lan.clear()
            await state.finish()
        else:
            await query.message.answer("Произошла ошибка. Повторите попытку:")
            os.remove(f"{query.from_user.id}.docx")
            os.remove(f"image_{query.from_user.id}.jpg")
            os.remove(f"handlers/users/image_{query.from_user.id}.jpg")
            know_lan.clear()
            await state.finish()


@dp.callback_query_handler(lambda query: query.data == "❌ Bekor qilish" or query.data == "❌ Отменить",
                           state=Full_info_user.comfirm)
async def get_name(query: types.CallbackQuery, state: FSMContext):
    language = await get_language(query.from_user.id)
    await query.message.answer(cancel_message_txt[language],
                               reply_markup=await get_markup_default(language, reg_btn_txt))
    os.remove(f"image_{query.from_user.id}.jpg")
    os.remove(f"handlers/users/image_{query.from_user.id}.jpg")
    know_lan.clear()
    await state.finish()
